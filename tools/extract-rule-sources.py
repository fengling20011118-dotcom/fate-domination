from pathlib import Path
from docx import Document
from openpyxl import load_workbook


PROJECT_DIR = Path(__file__).resolve().parent.parent
SOURCE_DIR = PROJECT_DIR.parent
OUTPUT_DIR = PROJECT_DIR / "docs" / "rule-sources"


def clean(value):
    return str(value or "").replace("\xa0", " ").strip()


def extract_docx(source: Path, output: Path):
    document = Document(source)
    lines = [f"# {source.stem}", ""]
    for paragraph in document.paragraphs:
        text = clean(paragraph.text)
        if text:
            lines.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        lines.extend(["", f"## 表格 {table_index}", ""])
        for row in table.rows:
            cells = [clean(cell.text).replace("\n", " / ") for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    output.write_text("\n".join(lines) + "\n", encoding="utf-8")


def extract_xlsx(source: Path, output: Path):
    workbook = load_workbook(source, data_only=False, read_only=True)
    lines = [f"# {source.stem}", ""]
    for worksheet in workbook.worksheets:
        lines.extend([f"## {worksheet.title}", ""])
        for row in worksheet.iter_rows():
            cells = [clean(cell.value).replace("\n", " / ") for cell in row]
            while cells and not cells[-1]:
                cells.pop()
            if any(cells):
                lines.append(" | ".join(cells))
        lines.append("")
    output.write_text("\n".join(lines), encoding="utf-8")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    extract_docx(
        SOURCE_DIR / "Fate_Domination 基础规则【墨水修订1版】.docx",
        OUTPUT_DIR / "base-rules.md",
    )
    extract_docx(
        SOURCE_DIR / "Fate_Domination FQA.docx",
        OUTPUT_DIR / "fqa.md",
    )
    extract_xlsx(
        SOURCE_DIR / "Fate-桌游问题解答和规则说明.xlsx",
        OUTPUT_DIR / "qa-workbook.md",
    )
    (OUTPUT_DIR / "turn-flow-and-keywords.md").write_text(
        (SOURCE_DIR / "玩家回合流程和关键词.txt").read_text(encoding="utf-8"),
        encoding="utf-8",
    )
    (OUTPUT_DIR / "3x-rules.md").write_text(
        (SOURCE_DIR / "3X模式规则.txt").read_text(encoding="utf-8"),
        encoding="utf-8",
    )
    print(f"规则来源已提取到：{OUTPUT_DIR}")


if __name__ == "__main__":
    main()
